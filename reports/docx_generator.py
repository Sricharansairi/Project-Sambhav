from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def generate_docx(data, output_path):
    """
    Professional DOCX generator for Project Sambhav.
    """
    doc = Document()
    
    # Header
    title = doc.add_heading(data.get("project_name", "Project Sambhav"), 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph(f"{data.get('subtitle', '')} — {data.get('version', 'v1.1')}")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph("_" * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Summary
    doc.add_heading("1. Prediction Summary", level=1)
    p = doc.add_paragraph()
    p.add_run("Prediction ID: ").bold = True
    p.add_run(data.get("prediction_id", "N/A") + "\t")
    p.add_run("Domain: ").bold = True
    p.add_run(data.get("domain", "N/A") + "\n")
    p.add_run("Generated At: ").bold = True
    p.add_run(data.get("generated_at", "N/A") + "\t")
    p.add_run("Mode: ").bold = True
    p.add_run(data.get("mode", "guided").title())

    # Probabilistic Analysis
    doc.add_heading("2. Probabilistic Inference Layers", level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Layer'
    hdr_cells[1].text = 'Probability'
    hdr_cells[2].text = 'Contribution'
    hdr_cells[3].text = 'Status'
    
    layers = [
        ["ML Prediction Layer", data.get("ml_probability", "N/A"), "65%", "CALIBRATED"],
        ["LLM Inference Layer", data.get("llm_probability", "N/A"), "35%", "STOCHASTIC"],
        ["Reconciled Final", data.get("reconciled_probability", "N/A"), "100%", "OPTIMIZED"]
    ]
    
    for layer in layers:
        row_cells = table.add_row().cells
        for i, val in enumerate(layer):
            row_cells[i].text = str(val)

    doc.add_paragraph(f"\nReliability Index: {data.get('reliability_index', 'N/A')}% ({data.get('warning_level', 'CLEAR')})")
    doc.add_paragraph(f"Agreement Gap: {data.get('agreement_gap', 'N/A')} | Confidence Tier: {data.get('confidence_tier', 'N/A')}")

    # Outcomes
    doc.add_heading("3. Detailed Outcomes", level=1)
    outcomes = data.get("outcomes", [])
    if outcomes:
        otable = doc.add_table(rows=1, cols=2)
        otable.style = 'Table Grid'
        otable.rows[0].cells[0].text = "Outcome Label"
        otable.rows[0].cells[1].text = "Probability Score"
        for o in outcomes:
            row = otable.add_row().cells
            row[0].text = o.get("label", "")
            row[1].text = o.get("probability", "")
    else:
        doc.add_paragraph("No specific outcomes generated.")

    # Parameters
    doc.add_heading("4. Input Parameters", level=1)
    params = data.get("parameters", {})
    if params:
        ptable = doc.add_table(rows=1, cols=2)
        ptable.style = 'Table Grid'
        ptable.rows[0].cells[0].text = "Parameter"
        ptable.rows[0].cells[1].text = "Value"
        for k, v in params.items():
            row = ptable.add_row().cells
            row[0].text = k.replace("_", " ").title()
            row[1].text = str(v)

    # SHAP
    doc.add_heading("5. Feature Contribution (SHAP)", level=1)
    shaps = data.get("shap_values", {})
    if shaps:
        stable = doc.add_table(rows=1, cols=3)
        stable.style = 'Table Grid'
        stable.rows[0].cells[0].text = "Feature"
        stable.rows[0].cells[1].text = "Impact Score"
        stable.rows[0].cells[2].text = "Direction"
        for k, v in shaps.items():
            row = stable.add_row().cells
            row[0].text = k.replace("_", " ").title()
            row[1].text = f"{v:+.4f}"
            row[2].text = "Positive (+)" if v > 0 else "Negative (-)"

    # Audit
    doc.add_heading("6. Audit & Safety Compliance", level=1)
    audit = data.get("audit", {})
    doc.add_paragraph(f"Overall Status: {audit.get('overall_status', 'PASSED')}")
    flags = audit.get("flags", [])
    if flags:
        for f in flags:
            doc.add_paragraph(f"• [{f.get('code', 'INFO')}] {f.get('message', '')} ({f.get('severity', 'INFO')})", style='List Bullet')
    else:
        doc.add_paragraph("• All safety and integrity checks passed.", style='List Bullet')

    # Footer
    doc.add_paragraph("\n" + "_" * 50)
    doc.add_paragraph(data.get("disclaimer", ""), style='Caption')
    doc.add_paragraph("© 2026 Sricharan Sairi. Generated via Project Sambhav Inference Engine.", style='Caption')

    doc.save(output_path)
