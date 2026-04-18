"""
api/endpoints/export.py — Full export engine for Project Sambhav.
Supports: PDF, Word, Excel, JSON, CSV, XML, PNG (chart), API Link.
"""
import os, json, io, logging
from fastapi import APIRouter, HTTPException, Depends
from fastapi.responses import StreamingResponse, JSONResponse
from pydantic import BaseModel, Field
from typing import Optional
from sqlalchemy.orm import Session
from db.models import get_db
from api.endpoints.auth import get_current_user

router = APIRouter()
logger = logging.getLogger(__name__)

EXPORTS_DIR = os.getenv("EXPORTS_DIR", "/tmp/sambhav_exports")
try:
    os.makedirs(EXPORTS_DIR, exist_ok=True)
except PermissionError:
    # Fallback for restricted environments like HuggingFace Spaces
    EXPORTS_DIR = "/tmp"
    logger.warning(f"Permission denied for exports directory. Falling back to {EXPORTS_DIR}")

BASE_URL = os.getenv("FASTAPI_URL", "http://localhost:8000")


def _fmt_pct(v, default: str = "—") -> str:
    """Safely format a 0-1 float as a percentage string. Returns default if value is None/invalid."""
    try:
        if v is None:
            return default
        return f"{float(v) * 100:.1f}%"
    except (TypeError, ValueError):
        return default


def _fmt_pct0(v, default: str = "—") -> str:
    """Safely format a 0-1 float as a rounded percentage (no decimal). Returns default if None/invalid."""
    try:
        if v is None:
            return default
        return f"{float(v) * 100:.0f}%"
    except (TypeError, ValueError):
        return default


def _fmt_float(v, fmt: str = "+.4f", default: str = "0.0000") -> str:
    """Safely format a float. Returns default if value is None/invalid."""
    try:
        if v is None:
            return default
        return format(float(v), fmt)
    except (TypeError, ValueError):
        return default


def _get_outcomes_list(data: dict) -> list:
    """Normalize outcomes to a list of (label, probability_pct) tuples regardless of source format."""
    # Prefer outcomes_list (set by frontend with full outcome objects)
    outcomes_list = data.get("outcomes_list", [])
    if outcomes_list and isinstance(outcomes_list, list):
        result = []
        for o in outcomes_list:
            if isinstance(o, dict):
                label = o.get("outcome", str(o))
                prob_raw = o.get("probability", 0)
                # probability may be 0-100 (frontend) or 0-1 (backend dict)
                prob_pct = f"{float(prob_raw):.1f}%" if float(prob_raw) > 1 else f"{float(prob_raw)*100:.1f}%"
                reasoning = o.get("reasoning", "")
                result.append((label, prob_pct, reasoning))
        if result:
            return result

    # Fallback: outcomes dict from prediction result
    outcomes = data.get("outcomes", {})
    if isinstance(outcomes, dict):
        return [(label, _fmt_pct(prob), "") for label, prob in outcomes.items()]
    return []


class ExportRequest(BaseModel):
    prediction_id: Optional[str] = Field(None, example="SMB-2026-00001")
    domain:        Optional[str] = None
    parameters:    Optional[dict]= {}
    result:        Optional[dict]= {}   # Full prediction result dict
    question:      Optional[str] = None
    format:        str           = Field("json", example="pdf")


def _get_prediction_data(req: ExportRequest, db: Session) -> dict:
    """Retrieve prediction data — from DB if prediction_id, else from req.result."""
    if req.prediction_id:
        try:
            from db.database import get_prediction
            pred = get_prediction(db, req.prediction_id)
            if pred:
                return pred
        except Exception as e:
            logger.warning(f"DB fetch failed: {e}")
    # Fall back to inline result
    return {
        "prediction_id": req.prediction_id or "SMB-EXPORT",
        "domain":        req.domain or "unknown",
        "question":      req.question or "",
        "parameters":    req.parameters or {},
        **(req.result or {}),
    }


# ── POST /export/json ─────────────────────────────────────────
@router.post("/json")
async def export_json(req: ExportRequest, db: Session = Depends(get_db)):
    data = _get_prediction_data(req, db)
    content = json.dumps(data, indent=2, default=str)
    return StreamingResponse(
        io.BytesIO(content.encode()),
        media_type="application/json",
        headers={"Content-Disposition": f"attachment; filename=sambhav_{data.get('prediction_id','export')}.json"}
    )


# ── POST /export/csv ──────────────────────────────────────────
@router.post("/csv")
async def export_csv(req: ExportRequest, db: Session = Depends(get_db)):
    import csv
    data = _get_prediction_data(req, db)
    output = io.StringIO()
    writer = csv.writer(output)

    headers = [
        "Prediction_ID", "Domain", "Question", "Execution_Mode", 
        "Final_Probability_Pct", "Reliability_Index_Pct", "System_Confidence",
        "ML_Probability_Pct", "LLM_Probability_Pct"
    ]
    
    # Collect parameters natively 
    params = data.get("raw_parameters") or data.get("parameters") or {}
    for k in params.keys():
        if not str(k).startswith("_"):
            headers.append(f"Param_{str(k)}")
            
    # Collect outcomes iteratively
    outcome_items = _get_outcomes_list(data)
    for i in range(len(outcome_items)):
        headers.extend([f"Outcome_{i+1}_Label", f"Outcome_{i+1}_ProbPct", f"Outcome_{i+1}_Reasoning"])
        
    writer.writerow(headers)

    row = [
        data.get("prediction_id", "N/A"),
        str(data.get("domain", "")).upper(),
        data.get("question", ""),
        data.get("mode", "guided"),
        round(data.get("final_probability", 0) * 100, 2),
        round(data.get("reliability_index", 0) * 100, 2),
        data.get("confidence_tier", ""),
        round(data.get("ml_probability", 0) * 100, 2) if data.get("ml_probability") else "",
        round(data.get("llm_probability", 0) * 100, 2) if data.get("llm_probability") else ""
    ]
    
    for k in params.keys():
        if not str(k).startswith("_"):
            row.append(str(params.get(k, "")))
            
    for item in outcome_items:
        row.extend([item[0], item[1].strip('%'), item[2]])

    writer.writerow(row)

    content = output.getvalue().encode()
    return StreamingResponse(
        io.BytesIO(content),
        media_type="text/csv",
        headers={"Content-Disposition": f"attachment; filename=Sambhav_Dataset_{data.get('prediction_id','export')}.csv"}
    )


# ── POST /export/xml ──────────────────────────────────────────
@router.post("/xml")
async def export_xml(req: ExportRequest, db: Session = Depends(get_db)):
    data = _get_prediction_data(req, db)

    def _val(v):
        return str(v).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

    lines = ['<?xml version="1.0" encoding="UTF-8"?>', "<SambhavPrediction>"]
    lines.append(f"  <PredictionId>{_val(data.get('prediction_id',''))}</PredictionId>")
    lines.append(f"  <Domain>{_val(data.get('domain',''))}</Domain>")
    lines.append(f"  <Question>{_val(data.get('question',''))}</Question>")
    lines.append(f"  <FinalProbability>{_val(_fmt_pct(data.get('final_probability')))}</FinalProbability>")
    lines.append(f"  <MLProbabilityLayer>{_val(_fmt_pct(data.get('ml_probability')))}</MLProbabilityLayer>")
    lines.append(f"  <LLMInferenceLayer>{_val(_fmt_pct(data.get('llm_probability')))}</LLMInferenceLayer>")
    lines.append(f"  <ReliabilityIndex>{_val(_fmt_pct(data.get('reliability_index')))}</ReliabilityIndex>")
    lines.append(f"  <ConfidenceTier>{_val(data.get('confidence_tier',''))}</ConfidenceTier>")
    lines.append(f"  <AgreementGap>{_val(_fmt_pct(data.get('gap')))}</AgreementGap>")
    lines.append(f"  <WarningLevel>{_val(data.get('warning_level','CLEAR'))}</WarningLevel>")
    lines.append(f"  <Mode>{_val(data.get('mode',''))}</Mode>")
    
    lines.append("  <Outcomes>")
    outcomes = data.get("outcomes", {})
    if isinstance(outcomes, dict):
        for label, prob in outcomes.items():
            lines.append(f"    <Outcome label=\"{_val(label)}\">{_val(_fmt_pct(prob))}</Outcome>")
    lines.append("  </Outcomes>")

    lines.append("  <Parameters>")
    for k, v in (data.get("raw_parameters") or data.get("parameters") or {}).items():
        lines.append(f"    <{k}>{_val(v)}</{k}>")
    lines.append("  </Parameters>")
    lines.append("  <SHAPValues>")
    for k, v in (data.get("shap_values") or {}).items():
        lines.append(f"    <Feature name=\"{_val(k)}\" impact=\"{_val(_fmt_float(v))}\" />")
    lines.append("  </SHAPValues>")
    lines.append("</SambhavPrediction>")

    content = "\n".join(lines).encode()
    return StreamingResponse(
        io.BytesIO(content),
        media_type="application/xml",
        headers={"Content-Disposition": f"attachment; filename=sambhav_{data.get('prediction_id','export')}.xml"}
    )


# ── POST /export/pdf ──────────────────────────────────────────
@router.post("/pdf")
async def export_pdf(req: ExportRequest, db: Session = Depends(get_db)):
    try:
        from reportlab.lib.pagesizes import A4
        import datetime
        from reportlab.lib import colors
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable, PageBreak

        data = _get_prediction_data(req, db)
        buf = io.BytesIO()
        doc = SimpleDocTemplate(buf, pagesize=A4, topMargin=2*cm, bottomMargin=2*cm,
                                leftMargin=2*cm, rightMargin=2*cm)

        styles = getSampleStyleSheet()
        
        # Light Corporate Theme Styles
        title_style = ParagraphStyle("title", parent=styles["Heading1"],
                                     textColor=colors.HexColor("#0f172a"), fontSize=28, spaceAfter=14,
                                     fontName="Helvetica-Bold", alignment=1) # Center
        subtitle_style = ParagraphStyle("subtitle", parent=styles["Normal"],
                                        textColor=colors.HexColor("#475569"), fontSize=12, spaceAfter=40, alignment=1)
        meta_style = ParagraphStyle("meta", parent=styles["Normal"],
                                    textColor=colors.HexColor("#64748b"), fontSize=10, alignment=1, spaceAfter=4)
        h2_style = ParagraphStyle("h2", parent=styles["Heading2"],
                                  textColor=colors.HexColor("#0f172a"), fontSize=14, spaceBefore=20, spaceAfter=8,
                                  fontName="Helvetica-Bold")
        p_style = ParagraphStyle("p", parent=styles["Normal"],
                                 textColor=colors.HexColor("#334155"), fontSize=10, leading=14, spaceAfter=8)
        cell_style = ParagraphStyle(name='cell', fontSize=9, textColor=colors.HexColor("#1e293b"))

        TABLE_STYLE = TableStyle([
            ("BACKGROUND",    (0, 0), (-1, 0), colors.HexColor("#f8fafc")),
            ("TEXTCOLOR",     (0, 0), (-1, 0), colors.HexColor("#334155")),
            ("FONTNAME",      (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE",      (0, 0), (-1,-1), 9),
            ("TEXTCOLOR",     (0, 1), (-1,-1), colors.HexColor("#1e293b")),
            ("ROWBACKGROUNDS",(0, 1), (-1,-1), [colors.white, colors.HexColor("#f1f5f9")]),
            ("GRID",          (0, 0), (-1,-1), 0.5, colors.HexColor("#e2e8f0")),
            ("PADDING",       (0, 0), (-1,-1), 8),
            ("VALIGN",        (0, 0), (-1,-1), "MIDDLE"),
        ])

        story = []

        # ── Cover Page ─────────────────────────────────────────────
        story.append(Spacer(1, 4*cm))
        story.append(Paragraph("Project Sambhav", title_style))
        story.append(Paragraph("Advanced Predictive & Forensic Intelligence Report", subtitle_style))
        story.append(Spacer(1, 2*cm))
        story.append(Paragraph(f"<b>Prediction ID:</b> {data.get('prediction_id', 'N/A')}", meta_style))
        story.append(Paragraph(f"<b>Generated:</b> {datetime.datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}", meta_style))
        story.append(Paragraph(f"<b>Domain Engine:</b> {str(data.get('domain', '')).upper()}", meta_style))
        story.append(PageBreak())

        # ── Prediction Summary ───────────────────────────────────────
        story.append(Paragraph("Executive Summary", h2_style))
        fp = data.get("final_probability", 0)
        ri = data.get("reliability_index", 0)

        meta_rows = [
            ["Metric", "Evaluation"],
            ["Primary Question", Paragraph(str(data.get("question", "Inference Evaluation")), cell_style)],
            ["Execution Mode",   str(data.get("mode", "guided")).capitalize()],
            ["Final Probability",_fmt_pct(fp)],
            ["ML Algorithmic Layer",     _fmt_pct(data.get('ml_probability'))],
            ["LLM Heuristic Layer",      _fmt_pct(data.get('llm_probability'))],
            ["Reliability Index",_fmt_pct0(ri)],
            ["System Confidence",str(data.get("confidence_tier", ""))],
        ]
        meta_t = Table(meta_rows, colWidths=[5.5*cm, 11.5*cm])
        meta_t.setStyle(TABLE_STYLE)
        story.append(meta_t)

        # ── Multi-Outcome Predictions ────────────────────────────────
        outcome_items = _get_outcomes_list(data)
        if outcome_items:
            story.append(Paragraph("Alternative Trajectory Models", h2_style))
            o_rows = [["Scenario Outcome", "Prob", "Heuristic Reasoning"]] + [
                [label, prob_pct, Paragraph(str(reasoning) if reasoning else "-", cell_style)]
                for label, prob_pct, reasoning in outcome_items
            ]
            o_t = Table(o_rows, colWidths=[4*cm, 1.5*cm, 11.5*cm])
            base_style = list(TABLE_STYLE.getCommands())
            base_style.append(("VALIGN", (2, 1), (-1, -1), "TOP"))
            o_t.setStyle(TableStyle(base_style))
            story.append(o_t)
            
        # ── Input Parameters ─────────────────────────────────────────
        params = data.get("raw_parameters") or data.get("parameters") or {}
        if params:
            story.append(Paragraph("Parameter Graph", h2_style))
            p_data = [[Paragraph(str(k).replace("_"," ").title(), cell_style), Paragraph(str(v), cell_style)] for k, v in params.items() if not str(k).startswith("_")]
            if p_data:
                p_table = Table([["Node", "State Value"]] + p_data, colWidths=[7*cm, 10*cm])
                p_table.setStyle(TABLE_STYLE)
                story.append(p_table)

        # ── Pragma Technical Appendix ──────────────────────────────
        if str(data.get("domain", "")).lower() == "pragma":
            story.append(PageBreak())
            story.append(Paragraph("Appendix A: PRAGMA Forensic Methodology", h2_style))
            story.append(Paragraph("This report was generated utilizing the PRAGMA (Pattern Recognition and Applied Geopolitical/Mental Analysis) architecture. PRAGMA acts as a cognitive load auditor, searching for linguistic and thematic indicators of deception, evasion, and memory fabrication.", p_style))
            story.append(Paragraph("<b>1. Linguistic Distancing:</b> PRAGMA tracks sudden pronoun shifts (drops in 'I' usage) signaling a subconscious separation from the account.", p_style))
            story.append(Paragraph("<b>2. Temporal Disassociation:</b> Accidental tense-hopping (reverting to present tense when discussing past events) strongly correlates with fabricated event sequencing.", p_style))
            story.append(Paragraph("<b>3. Evasion & Over-Qualification:</b> Detection of buffer phrases ('to be honest') paired with missing categorical details indicates cognitive friction.", p_style))
            story.append(Paragraph("<b>4. Route Engine Adversarial Auditing:</b> PRAGMA employs a dual-agent layout where one matrix builds the narrative while a 'Devil\'s Advocate' explicitly tests it for plausibility mechanics.", p_style))
            story.append(Paragraph("<i>Disclaimer: PRAGMA provides probabilistic behavioral vectors. It is strictly advisory and cannot substitute certified polygraphic, forensic, or legal auditing.</i>", p_style))

        # ── Footer ──────────────────────────────────────────────────
        story.append(Spacer(1, 1*cm))
        story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#cbd5e1")))
        story.append(Spacer(1, 0.2*cm))
        story.append(Paragraph(
            "Sambhav inference models provide probabilistic estimations based on multi-variate topologies. Always utilize human expert calibration.",
            ParagraphStyle("disclaimer", parent=styles["Normal"], textColor=colors.HexColor("#94a3b8"), fontSize=7, alignment=1)
        ))

        doc.build(story)
        buf.seek(0)
        return StreamingResponse(
            buf,
            media_type="application/pdf",
            headers={"Content-Disposition": f"attachment; filename=Sambhav_Report_{data.get('prediction_id','export')}.pdf"}
        )
    except ImportError:
        raise HTTPException(status_code=500, detail="ReportLab not installed. Run: pip install reportlab")
    except Exception as e:
        logger.error(f"PDF export error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# ── POST /export/word ─────────────────────────────────────────
@router.post("/word")
async def export_word(req: ExportRequest, db: Session = Depends(get_db)):
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        data = _get_prediction_data(req, db)
        doc = Document()

        # ── Setup Formatting ──────────────────────────────────────
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(10)

        # ── Cover Page ─────────────────────────────────────────────
        import datetime
        doc.add_paragraph().add_run().add_break()
        doc.add_paragraph().add_run().add_break()
        title = doc.add_heading("Project Sambhav", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub = doc.add_paragraph("Advanced Predictive & Forensic Intelligence Report")
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph().add_run().add_break()
        doc.add_paragraph().add_run().add_break()
        
        meta_pg = doc.add_paragraph()
        meta_pg.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta_pg.add_run("Prediction ID: ").bold = True
        meta_pg.add_run(f"{data.get('prediction_id', 'N/A')}\n")
        meta_pg.add_run("Generated: ").bold = True
        meta_pg.add_run(f"{datetime.datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}\n")
        meta_pg.add_run("Domain Engine: ").bold = True
        meta_pg.add_run(f"{str(data.get('domain', '')).upper()}")
        
        doc.add_page_break()

        # ── Executive Summary ──────────────────────────────────────
        doc.add_heading("Executive Summary", level=1)
        fp = data.get("final_probability", 0)
        rows_data = [
            ("Primary Question", str(data.get("question", "Inference Evaluation"))),
            ("Execution Mode",   str(data.get("mode", "guided")).capitalize()),
            ("Final Probability", _fmt_pct(fp)),
            ("ML Algorithmic Layer",     _fmt_pct(data.get('ml_probability'))),
            ("LLM Heuristic Layer",      _fmt_pct(data.get('llm_probability'))),
            ("Reliability Index", _fmt_pct0(data.get('reliability_index'))),
            ("System Confidence",  str(data.get("confidence_tier", ""))),
        ]
        table = doc.add_table(rows=len(rows_data), cols=2)
        table.style = 'Light Shading Accent 1'
        for i, (label, value) in enumerate(rows_data):
            table.rows[i].cells[0].text = label
            table.rows[i].cells[1].text = value
            table.rows[i].cells[0].paragraphs[0].runs[0].bold = True

        doc.add_paragraph()

        # ── Alternative Trajectory Models ──────────────────────────
        outcome_items = _get_outcomes_list(data)
        if outcome_items:
            doc.add_heading("Alternative Trajectory Models", level=1)
            o_table = doc.add_table(rows=1 + len(outcome_items), cols=3)
            o_table.style = 'Light Shading Accent 1'
            hdr = o_table.rows[0].cells
            hdr[0].text, hdr[1].text, hdr[2].text = "Scenario Outcome", "Prob", "Heuristic Reasoning"
            for cell in hdr:
                cell.paragraphs[0].runs[0].bold = True
            for i, (label, prob_pct, reasoning) in enumerate(outcome_items, 1):
                row = o_table.rows[i].cells
                row[0].text = label
                row[1].text = prob_pct
                row[2].text = str(reasoning) if reasoning else "-"

        doc.add_paragraph()

        # ── Parameter Graph ────────────────────────────────────────
        params = data.get("raw_parameters") or data.get("parameters") or {}
        if params:
            doc.add_heading("Parameter Graph", level=1)
            p_data = [(str(k).replace("_"," ").title(), str(v)) for k, v in params.items() if not str(k).startswith("_")]
            
            p_table = doc.add_table(rows=1 + len(p_data), cols=2)
            p_table.style = 'Light Shading Accent 1'
            h = p_table.rows[0].cells
            h[0].text, h[1].text = "Node", "State Value"
            for cell in h: cell.paragraphs[0].runs[0].bold = True
                
            for i, (label, value) in enumerate(p_data, 1):
                p_table.rows[i].cells[0].text = label
                p_table.rows[i].cells[1].text = str(value)

        # ── Pragma Technical Appendix ──────────────────────────────
        if str(data.get("domain", "")).lower() == "pragma":
            doc.add_page_break()
            doc.add_heading("Appendix A: PRAGMA Forensic Methodology", level=1)
            doc.add_paragraph("This report was generated utilizing the PRAGMA (Pattern Recognition and Applied Geopolitical/Mental Analysis) architecture. PRAGMA acts as a cognitive load auditor, searching for linguistic and thematic indicators of deception, evasion, and memory fabrication.")
            
            p1 = doc.add_paragraph()
            p1.add_run("1. Linguistic Distancing: ").bold = True
            p1.add_run("PRAGMA tracks sudden pronoun shifts (drops in 'I' usage) signaling a subconscious separation from the account.")
            
            p2 = doc.add_paragraph()
            p2.add_run("2. Temporal Disassociation: ").bold = True
            p2.add_run("Accidental tense-hopping (reverting to present tense when discussing past events) strongly correlates with fabricated event sequencing.")
            
            p3 = doc.add_paragraph()
            p3.add_run("3. Evasion & Over-Qualification: ").bold = True
            p3.add_run("Detection of buffer phrases ('to be honest') paired with missing categorical details indicates cognitive friction.")
            
            p4 = doc.add_paragraph()
            p4.add_run("4. Route Engine Adversarial Auditing: ").bold = True
            p4.add_run("PRAGMA employs a dual-agent layout where one matrix builds the narrative while a 'Devil\'s Advocate' explicitly tests it for plausibility mechanics.")
            
            disclaimer = doc.add_paragraph("Disclaimer: PRAGMA provides probabilistic behavioral vectors. It is strictly advisory and cannot substitute certified polygraphic, forensic, or legal auditing.")
            disclaimer.style.font.italic = True

        # ── File Output ────────────────────────────────────────────
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return StreamingResponse(
            buf,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename=Sambhav_Report_{data.get('prediction_id','export')}.docx"}
        )
    except ImportError:
        raise HTTPException(status_code=500, detail="python-docx not installed. Run: pip install python-docx")
    except Exception as e:
        logger.error(f"Word export error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# ── POST /export/excel ────────────────────────────────────────
@router.post("/excel")
async def export_excel(req: ExportRequest, db: Session = Depends(get_db)):
    try:
        import openpyxl
        from openpyxl.styles import Font, PatternFill, Alignment

        data = _get_prediction_data(req, db)
        wb = openpyxl.Workbook()

        # Sheet 1: Summary
        ws = wb.active
        ws.title = "Summary"
        header_fill = PatternFill("solid", fgColor="1f2937")
        header_font = Font(color="ffffff", bold=True)
        val_font    = Font(color="111827")

        rows = [
            ("Field", "Value"),
            ("Prediction ID",   data.get("prediction_id", "")),
            ("Domain",          str(data.get("domain", "")).upper()),
            ("Question",        data.get("question", "")),
            ("Final Probability", _fmt_pct(data.get('final_probability'))),
            ("ML Probability",  _fmt_pct(data.get('ml_probability'))),
            ("LLM Probability", _fmt_pct(data.get('llm_probability'))),
            ("Reliability Index", _fmt_pct0(data.get('reliability_index'))),
            ("Confidence Tier", data.get("confidence_tier", "")),
            ("Mode",            data.get("mode", "")),
        ]
        for r_idx, (field, value) in enumerate(rows, 1):
            ws.cell(r_idx, 1, field).font  = header_font if r_idx == 1 else val_font
            ws.cell(r_idx, 2, str(value)).font = val_font
            if r_idx == 1:
                ws.cell(r_idx, 1).fill = header_fill
                ws.cell(r_idx, 2).fill = header_fill
        ws.column_dimensions["A"].width = 22
        ws.column_dimensions["B"].width = 40

        # Sheet 2: Parameters
        ws2 = wb.create_sheet("Parameters")
        ws2.cell(1, 1, "Parameter").font = header_font
        ws2.cell(1, 1).fill = header_fill
        ws2.cell(1, 2, "Value").font     = header_font
        ws2.cell(1, 2).fill = header_fill
        params = data.get("raw_parameters") or data.get("parameters") or {}
        for i, (k, v) in enumerate(params.items(), 2):
            if not str(k).startswith("_"):
                ws2.cell(i, 1, str(k)).font = val_font
                ws2.cell(i, 2, str(v)).font = val_font
        ws2.column_dimensions["A"].width = 25
        ws2.column_dimensions["B"].width = 30

        # Sheet 3: Multi-Outcome Predictions
        ws4 = wb.create_sheet("Outcomes")
        for col_idx, col_header in enumerate(["Outcome", "Probability", "Reasoning"], 1):
            cell = ws4.cell(1, col_idx, col_header)
            cell.font = header_font
            cell.fill = header_fill
        for i, (label, prob_pct, reasoning) in enumerate(_get_outcomes_list(data), 2):
            ws4.cell(i, 1, label).font    = val_font
            ws4.cell(i, 2, prob_pct).font = val_font
            ws4.cell(i, 3, str(reasoning)).font = val_font
        ws4.column_dimensions["A"].width = 35
        ws4.column_dimensions["B"].width = 15
        ws4.column_dimensions["C"].width = 60

        # Sheet 4: Flat Data Matrix (For ML Ingest)
        ws5 = wb.create_sheet("Flat Data Matrix")
        headers = [
            "Prediction_ID", "Domain", "Question", "Execution_Mode", 
            "Final_Probability_Pct", "Reliability_Index_Pct", "System_Confidence",
            "ML_Probability_Pct", "LLM_Probability_Pct"
        ]
        
        for k in params.keys():
            if not str(k).startswith("_"): headers.append(f"Param_{str(k)}")
                
        outcome_items = _get_outcomes_list(data)
        for i in range(len(outcome_items)):
            headers.extend([f"Outcome_{i+1}_Label", f"Outcome_{i+1}_ProbPct", f"Outcome_{i+1}_Reasoning"])
            
        for col_idx, h in enumerate(headers, 1):
            cell = ws5.cell(1, col_idx, h)
            cell.font = header_font
            cell.fill = header_fill

        row = [
            data.get("prediction_id", "N/A"),
            str(data.get("domain", "")).upper(),
            data.get("question", ""),
            data.get("mode", "guided"),
            round(data.get("final_probability", 0) * 100, 2),
            round(data.get("reliability_index", 0) * 100, 2),
            data.get("confidence_tier", ""),
            round(data.get("ml_probability", 0) * 100, 2) if data.get("ml_probability") else "",
            round(data.get("llm_probability", 0) * 100, 2) if data.get("llm_probability") else ""
        ]
        for k in params.keys():
            if not str(k).startswith("_"): row.append(str(params.get(k, "")))
        for item in outcome_items:
            row.extend([item[0], item[1].strip('%'), item[2]])

        for col_idx, val in enumerate(row, 1):
            ws5.cell(2, col_idx, str(val)).font = val_font

        buf = io.BytesIO()
        wb.save(buf)
        buf.seek(0)
        return StreamingResponse(
            buf,
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            headers={"Content-Disposition": f"attachment; filename=Sambhav_Report_{data.get('prediction_id','export')}.xlsx"}
        )
    except ImportError:
        raise HTTPException(status_code=500, detail="openpyxl not installed. Run: pip install openpyxl")
    except Exception as e:
        logger.error(f"Excel export error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# ── POST /export/png ──────────────────────────────────────────
@router.post("/png")
async def export_png(req: ExportRequest, db: Session = Depends(get_db)):
    try:
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt
        import matplotlib.patches as mpatches
        import numpy as np

        data = _get_prediction_data(req, db)
        shap = data.get("shap_values", {})

        fig, axes = plt.subplots(1, 2, figsize=(14, 6), facecolor="#08080a")

        # Left: probability gauge
        ax1 = axes[0]
        ax1.set_facecolor("#08080a")
        fp = float(data.get("final_probability") or 0.5)
        theta = np.linspace(np.pi, 0, 200)
        ax1.plot(np.cos(theta), np.sin(theta), color="#1a1a1a", linewidth=20, solid_capstyle="round")
        theta_fill = np.linspace(np.pi, np.pi - fp * np.pi, 200)
        color = "#c0c0c0" if fp >= 0.5 else "#ffb7c5"
        ax1.plot(np.cos(theta_fill), np.sin(theta_fill), color=color, linewidth=20, solid_capstyle="round")
        ax1.text(0, -0.1, f"{fp*100:.1f}%", ha="center", va="center", fontsize=28, color="#e8e8e8", fontweight="bold")
        ax1.text(0, -0.35, data.get("domain","").upper(), ha="center", va="center", fontsize=11, color="#a0a0a0")
        ax1.text(0, -0.55, f"Reliability: {_fmt_pct0(data.get('reliability_index'))}", ha="center", va="center", fontsize=9, color="#666666")
        ax1.set_xlim(-1.3, 1.3)
        ax1.set_ylim(-0.8, 1.3)
        ax1.axis("off")
        ax1.set_title("Prediction Probability", color="#c0c0c0", fontsize=12, pad=10)

        # Right: SHAP bar chart
        ax2 = axes[1]
        ax2.set_facecolor("#08080a")
        if shap:
            items = sorted(shap.items(), key=lambda x: abs(float(x[1])) if isinstance(x[1],(int,float)) else 0, reverse=True)[:8]
            features = [str(k)[:20] for k, _ in items]
            values   = [float(v) if isinstance(v,(int,float)) else 0 for _, v in items]
            bar_colors = ["#c0c0c0" if v >= 0 else "#ffb7c5" for v in values]
            bars = ax2.barh(features, values, color=bar_colors, alpha=0.8)
            ax2.axvline(0, color="#333333", linewidth=0.8)
            ax2.tick_params(colors="#a0a0a0", labelsize=9)
            ax2.set_xlabel("SHAP Contribution", color="#a0a0a0", fontsize=9)
            for spine in ax2.spines.values():
                spine.set_edgecolor("#222222")
        ax2.set_title("Feature Contributions (SHAP)", color="#c0c0c0", fontsize=12, pad=10)

        fig.suptitle("Project Sambhav — Prediction Report", color="#e8e8e8", fontsize=14, y=1.02)
        plt.tight_layout()

        buf = io.BytesIO()
        plt.savefig(buf, format="png", dpi=150, bbox_inches="tight",
                    facecolor="#08080a", edgecolor="none")
        plt.close(fig)
        buf.seek(0)
        return StreamingResponse(
            buf,
            media_type="image/png",
            headers={"Content-Disposition": f"attachment; filename=sambhav_{data.get('prediction_id','export')}.png"}
        )
    except ImportError:
        raise HTTPException(status_code=500, detail="matplotlib not installed. Run: pip install matplotlib")
    except Exception as e:
        logger.error(f"PNG export error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# ── POST /export/api-link ─────────────────────────────────────
@router.post("/api-link")
async def export_api_link(req: ExportRequest):
    """Returns a shareable read-only API URL for this prediction."""
    pid = req.prediction_id or "demo"
    return JSONResponse({
        "success":      True,
        "api_url":      f"{BASE_URL}/history/{pid}",
        "json_url":     f"{BASE_URL}/export/json",
        "expires_in":   "72 hours",
        "instructions": "POST to api_url with your token to retrieve full prediction data.",
        "curl_example": f"curl -X POST {BASE_URL}/export/json -H 'Content-Type: application/json' -d '{{\"prediction_id\": \"{pid}\"}}'",
    })


# ── Generic dispatch ──────────────────────────────────────────
@router.post("")
@router.post("/{format}")
async def export_dispatch(
    req: ExportRequest,
    format: Optional[str] = None,
    db: Session = Depends(get_db)
):
    """Single endpoint that dispatches to format-specific handler."""
    # Prioritize path param over body param
    fmt = (format or req.format or "json").lower()
    dispatch = {
        "json":  export_json,
        "csv":   export_csv,
        "xml":   export_xml,
        "pdf":   export_pdf,
        "word":  export_word,
        "excel": export_excel,
        "png":   export_png,
        "api":   export_api_link,
    }
    handler = dispatch.get(fmt)
    if not handler:
        raise HTTPException(status_code=400, detail=f"Unsupported format: {fmt}. Choose from: {list(dispatch.keys())}")
    if fmt == "api":
        return await handler(req)
    return await handler(req, db)